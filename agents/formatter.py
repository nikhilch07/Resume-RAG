# formatter.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.schema import Resume

BLUE = RGBColor(0x1F, 0x4E, 0x79)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

def create_resume_docx(resume: Resume, output_path: str) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Name - 28pt heading
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.name)
    name_run.bold = True
    name_run.font.size = Pt(28)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(resume.contact)
    contact_run.font.size = Pt(11)
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER



    # Summary
    add_section_header(doc, "Summary")
    for line in resume.summary:
        # doc.add_paragraph(line, style="List Bullet")
        doc.add_paragraph(" ".join(resume.summary))   # join list into one paragraph, no bullet style

    # Skills
    add_section_header(doc, "Skills")
    for line in resume.skills:
        doc.add_paragraph(line, style="List Bullet")

    # Career Highlights
    add_section_header(doc, "Career Highlights")
    for line in resume.career_highlights:
        doc.add_paragraph(line, style="List Bullet")

    # Experience
    add_section_header(doc, "Experience")
    for job in resume.experience:
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        title_run = title_p.add_run(job.job_title.upper())
        title_run.bold = True
        title_run.font.size = Pt(11)

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(0)
        meta_run = meta_p.add_run(f"{job.company} | {job.location} | {job.tenure}")
        meta_run.bold = True
        meta_run.font.size = Pt(11)

        for bullet in job.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    # Education
    add_section_header(doc, "Education")
    for line in resume.education:
        doc.add_paragraph(line)

    if resume.portfolio:
        add_section_header(doc, "Websites & Portfolios")
        doc.add_paragraph(resume.portfolio)

    doc.save(output_path)